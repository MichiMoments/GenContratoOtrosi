"""Generación determinística del documento de otrosí.

Sin dependencia de Streamlit: recibe el payload y devuelve Markdown o .docx, para
que un futuro modo masivo pueda reutilizar este módulo tal cual.
"""

import io
import re
import unicodedata
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, StrictUndefined

# PLACEHOLDER — confirmar con RRHH antes de usar en firma.
EMPLEADOR = {
    "razon_social": "UNIVERSIDAD DE LOS ANDES",
    "nit": "[NIT]",
    "representante": "[NOMBRE DEL REPRESENTANTE LEGAL]",
    "cargo_representante": "[CARGO DEL REPRESENTANTE]",
}

PLANTILLAS = {
    1: "1_incremento_salarial.md.j2",
    2: "2_cambio_cargo.md.j2",
    3: "3_cambio_lugar.md.j2",
    4: "4_cambio_jornada.md.j2",
    5: "5_renovacion_contrato.md.j2",
}

MESES = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)


def moneda(valor):
    """1500000 -> $1.500.000 (formato colombiano, sin depender de `locale`)."""
    if valor is None:
        return ""
    return "$" + f"{int(valor):,}".replace(",", ".")


def fecha_larga(valor):
    """date/datetime -> '3 de agosto de 2026'."""
    if valor is None:
        return ""
    return f"{valor.day} de {MESES[valor.month - 1]} de {valor.year}"


def fecha_hora_larga(valor):
    """datetime -> '3 de agosto de 2026 a las 10:47'."""
    if valor is None:
        return ""
    return f"{fecha_larga(valor)} a las {valor:%H:%M}"


_env = Environment(
    loader=FileSystemLoader(Path(__file__).parent / "plantillas"),
    undefined=StrictUndefined,
    trim_blocks=True,
    lstrip_blocks=True,
    keep_trailing_newline=True,
)
_env.filters["moneda"] = moneda
_env.filters["fecha_larga"] = fecha_larga
_env.filters["fecha_hora_larga"] = fecha_hora_larga


def render_markdown(payload):
    """Renderiza el otrosí como Markdown a partir del payload del formulario."""
    plantilla = _env.get_template(PLANTILLAS[payload["tipo_id"]])
    return plantilla.render(
        empleador=EMPLEADOR,
        tipo=payload["tipo"],
        generales=payload["generales"],
        detalle=payload["detalle"],
    )


_NEGRITA = re.compile(r"\*\*(.+?)\*\*")


def _escribir_parrafo(doc, texto, estilo=None):
    """Agrega un párrafo partiendo `**negrita**` en runs separados."""
    parrafo = doc.add_paragraph(style=estilo)
    for i, fragmento in enumerate(_NEGRITA.split(texto)):
        if not fragmento:
            continue
        run = parrafo.add_run(fragmento)
        run.bold = i % 2 == 1  # los índices impares vienen de dentro de los **
    return parrafo


def markdown_a_docx(md):
    """Convierte el Markdown de las plantillas a un .docx y devuelve los bytes.

    Soporta solo el subconjunto que usan las plantillas: encabezados `#`/`##`,
    viñetas `- `, negrita `**...**`, y párrafos separados por línea en blanco.
    Las plantillas deben limitarse a eso.
    """
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    for linea in md.splitlines():
        linea = linea.strip()
        if not linea:
            continue
        if linea.startswith("## "):
            doc.add_heading(linea[3:].strip(), level=2)
        elif linea.startswith("# "):
            doc.add_heading(linea[2:].strip(), level=1)
        elif linea.startswith("- "):
            _escribir_parrafo(doc, linea[2:].strip(), estilo="List Bullet")
        else:
            _escribir_parrafo(doc, linea)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _slug(texto):
    """'David Pérez' -> 'david_perez'."""
    sin_tildes = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "_", sin_tildes.lower()).strip("_")


def nombre_archivo(payload):
    """Nombre del .docx: otrosi_2_david_perez_20260803.docx."""
    generales = payload["generales"]
    fecha = generales["fecha_otrosi"]
    if isinstance(fecha, (datetime, date)):
        fecha = f"{fecha:%Y%m%d}"
    return f"otrosi_{payload['tipo_id']}_{_slug(generales['nombre_empleado'])}_{fecha}.docx"
